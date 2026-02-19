"""Stage 6: Convert HTML reports to JSON and DOCX formats."""

import json
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _parse_html_structure(html_content: str, filename: str) -> dict:
    """Parse HTML content into a structured dictionary."""
    soup = BeautifulSoup(html_content, 'html.parser')
    data = {}

    title_tag = soup.find('h1')
    data['main_title'] = title_tag.get_text(separator='\n', strip=True) if title_tag else Path(filename).stem

    data['sections_data'] = {'introductory_content': []}
    sections_map = {}

    current_h2 = None
    current_h3 = None
    current_h4 = None

    def get_target():
        if current_h4 and current_h3 and current_h2:
            return sections_map[current_h2]['subsections'][current_h3]['sub_subsections'][current_h4]['content']
        elif current_h3 and current_h2:
            return sections_map[current_h2]['subsections'][current_h3]['content']
        elif current_h2:
            return sections_map[current_h2]['content']
        return data['sections_data']['introductory_content']

    container = soup.find('div', class_='report-container') or soup.body
    if not container:
        return None

    for tag in container.find_all(['h2', 'h3', 'h4', 'p', 'table', 'ul', 'ol', 'hr']):
        if tag.name == 'h2':
            current_h2 = tag.get_text(strip=True)
            sections_map[current_h2] = {'content': [], 'subsections': {}}
            current_h3 = None
            current_h4 = None
        elif tag.name == 'h3':
            if current_h2:
                current_h3 = tag.get_text(strip=True)
                sections_map[current_h2]['subsections'][current_h3] = {'content': [], 'sub_subsections': {}}
                current_h4 = None
            else:
                get_target().append({'type': 'heading', 'level': 3, 'text': tag.get_text(strip=True)})
        elif tag.name == 'h4':
            if current_h2 and current_h3:
                current_h4 = tag.get_text(strip=True)
                sections_map[current_h2]['subsections'][current_h3]['sub_subsections'][current_h4] = {'content': []}
            else:
                get_target().append({'type': 'heading', 'level': 4, 'text': tag.get_text(strip=True)})
        elif tag.name == 'p':
            text = tag.get_text(strip=True)
            if text:
                is_bold = bool(tag.find('strong') or tag.find('b'))
                item = {'type': 'paragraph', 'text': text, 'bold': is_bold}
                if tag.has_attr('class') and any(c in tag['class'] for c in
                    ['source-note', 'disclaimer', 'table-caption', 'bdo-header-footer']):
                    item['p_class'] = tag['class']
                get_target().append(item)
        elif tag.name in ['ul', 'ol']:
            items = [li.get_text(strip=True) for li in tag.find_all('li', recursive=False)
                     if li.get_text(strip=True)]
            if items:
                get_target().append({'type': tag.name, 'items': items})
        elif tag.name == 'table':
            caption = "Table"
            cap_tag = tag.find('caption')
            if cap_tag:
                caption = cap_tag.get_text(strip=True)
            table_data = {'headers': [], 'rows': []}
            thead = tag.find('thead')
            if thead:
                hr = thead.find('tr')
                if hr:
                    table_data['headers'] = [th.get_text(strip=True) for th in hr.find_all('th')]
            tbody = tag.find('tbody')
            if tbody:
                for row in tbody.find_all('tr'):
                    cols = [td.get_text(strip=True) for td in row.find_all('td')]
                    if cols:
                        table_data['rows'].append(cols)
            if not table_data['headers'] and table_data['rows'] and len(table_data['rows']) > 1:
                table_data['headers'] = table_data['rows'].pop(0)
            if table_data['headers'] or table_data['rows']:
                get_target().append({'type': 'table', 'caption': caption, 'data': table_data})
        elif tag.name == 'hr':
            get_target().append({'type': 'horizontal_rule'})

    data['sections_data']['sections_content'] = sections_map
    return data


def convert_html_to_json(html_content: str, filename: str,
                         output_dir: Path = None) -> dict:
    """Convert HTML content to structured JSON.

    Returns dict with 'json_data' and optionally 'output_path'.
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    data = {}
    title_tag = soup.find('h1')
    data['main_title'] = title_tag.get_text(separator='\n', strip=True) if title_tag else Path(filename).stem

    data['sections'] = []
    container = soup.find('div', class_='report-container') or soup.body or soup

    current_section = None
    current_subsection = None

    for tag in container.find_all(['h2', 'h3', 'h4', 'p', 'table', 'ul', 'ol']):
        if tag.name == 'h2':
            current_section = {'title': tag.get_text(strip=True), 'level': 2, 'content': [], 'subsections': []}
            data['sections'].append(current_section)
            current_subsection = None
        elif tag.name == 'h3' and current_section:
            current_subsection = {'title': tag.get_text(strip=True), 'level': 3, 'content': []}
            current_section['subsections'].append(current_subsection)
        elif tag.name == 'p':
            text = tag.get_text(strip=True)
            if text:
                target = current_subsection or current_section
                if target:
                    target['content'].append({'type': 'paragraph', 'text': text})
        elif tag.name in ['ul', 'ol']:
            items = [li.get_text(strip=True) for li in tag.find_all('li', recursive=False)]
            if items:
                target = current_subsection or current_section
                if target:
                    target['content'].append({'type': 'list', 'list_type': tag.name, 'items': items})
        elif tag.name == 'table':
            headers = []
            rows = []
            thead = tag.find('thead')
            if thead:
                hr = thead.find('tr')
                if hr:
                    headers = [th.get_text(strip=True) for th in hr.find_all('th')]
            tbody = tag.find('tbody')
            if tbody:
                for row in tbody.find_all('tr'):
                    cols = [td.get_text(strip=True) for td in row.find_all('td')]
                    if cols:
                        rows.append(cols)
            if headers or rows:
                target = current_subsection or current_section
                if target:
                    target['content'].append({'type': 'table', 'headers': headers, 'rows': rows})

    json_str = json.dumps(data, indent=4, ensure_ascii=False)

    output_path = None
    if output_dir:
        output_dir.mkdir(parents=True, exist_ok=True)
        stem = Path(filename).stem
        json_path = output_dir / (stem + ".json")
        json_path.write_text(json_str, encoding='utf-8')
        output_path = json_path

    return {"json_data": data, "json_str": json_str, "output_path": output_path,
            "filename": Path(filename).stem + ".json"}


def convert_html_to_docx(html_content: str, filename: str,
                         output_dir: Path = None) -> dict:
    """Convert HTML content to DOCX format.

    Returns dict with 'docx_bytes' and optionally 'output_path'.
    """
    structured = _parse_html_structure(html_content, filename)
    if not structured:
        raise RuntimeError(f"Could not parse HTML structure from {filename}")

    doc = Document()

    def add_items(parent, items):
        for item in items:
            if item['type'] == 'paragraph':
                if 'bdo-header-footer' in item.get('p_class', []) or \
                   'disclaimer' in item.get('p_class', []):
                    continue
                p = parent.add_paragraph()
                run = p.add_run(item['text'])
                if item.get('bold'):
                    run.bold = True
            elif item['type'] == 'heading':
                parent.add_heading(item['text'], level=item['level'])
            elif item['type'] in ['ul', 'ol']:
                style = 'ListBullet' if item['type'] == 'ul' else 'ListNumber'
                for li in item['items']:
                    parent.add_paragraph(li, style=style)
            elif item['type'] == 'table':
                tbl = item['data']
                if tbl['headers'] or tbl['rows']:
                    parent.add_paragraph(item.get('caption', 'Table'), style='Caption')
                    ncols = len(tbl['headers']) or (len(tbl['rows'][0]) if tbl['rows'] else 1)
                    ncols = max(ncols, 1)
                    doc_table = parent.add_table(rows=0, cols=ncols)
                    doc_table.style = 'TableGrid'
                    if tbl['headers']:
                        cells = doc_table.add_row().cells
                        for i, h in enumerate(tbl['headers']):
                            if i < ncols:
                                cells[i].text = h
                                cells[i].paragraphs[0].runs[0].font.bold = True
                    for row in tbl['rows']:
                        if len(row) == ncols:
                            cells = doc_table.add_row().cells
                            for i, c in enumerate(row):
                                cells[i].text = c
                    parent.add_paragraph()
            elif item['type'] == 'horizontal_rule':
                parent.add_paragraph("_________________________________________")

    # Title
    if structured.get('main_title'):
        for line in structured['main_title'].split('\n'):
            h = doc.add_heading(line.strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # Intro content
    add_items(doc, structured['sections_data']['introductory_content'])

    # Sections
    for h2_title, h2_data in structured['sections_data']['sections_content'].items():
        doc.add_heading(h2_title, level=1)
        add_items(doc, h2_data.get('content', []))
        for h3_title, h3_data in h2_data.get('subsections', {}).items():
            doc.add_heading(h3_title, level=2)
            add_items(doc, h3_data.get('content', []))
            for h4_title, h4_data in h3_data.get('sub_subsections', {}).items():
                doc.add_heading(h4_title, level=3)
                add_items(doc, h4_data.get('content', []))
        doc.add_paragraph()

    # Save to bytes
    import io
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    output_path = None
    if output_dir:
        output_dir.mkdir(parents=True, exist_ok=True)
        stem = Path(filename).stem
        docx_path = output_dir / (stem + ".docx")
        docx_path.write_bytes(docx_bytes)
        output_path = docx_path

    return {"docx_bytes": docx_bytes, "output_path": output_path,
            "filename": Path(filename).stem + ".docx"}


def convert_all_reports(html_items: list[dict] = None,
                        output_dir: Path = None,
                        log_callback=None) -> dict:
    """Convert a list of HTML items to JSON and DOCX.

    html_items: list of dicts with 'html_content' and 'filename' keys.
    Returns dict with 'json_files' and 'docx_files' lists.
    """
    def log(msg):
        if log_callback:
            log_callback(msg)

    if not html_items:
        html_items = []

    json_files = []
    docx_files = []

    for item in html_items:
        html_content = item["html_content"]
        filename = item["filename"]

        try:
            log(f"Converting {filename} to JSON...")
            result = convert_html_to_json(html_content, filename, output_dir)
            json_files.append(result)
        except Exception as e:
            log(f"JSON conversion failed for {filename}: {e}")

        try:
            log(f"Converting {filename} to DOCX...")
            result = convert_html_to_docx(html_content, filename, output_dir)
            docx_files.append(result)
        except Exception as e:
            log(f"DOCX conversion failed for {filename}: {e}")

    return {"json_files": json_files, "docx_files": docx_files}
